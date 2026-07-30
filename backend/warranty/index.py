import json
import os
import re
import hmac
import io
import urllib.request
from copy import deepcopy
from datetime import datetime

import psycopg2
import boto3
import docx


ORDERS_TABLE = 't_p288352_240fps_site_project.orders'
ITEMS_TABLE = 't_p288352_240fps_site_project.order_items'

TEMPLATE_URL = 'https://cdn.poehali.dev/projects/5376b460-4536-4f54-ba9a-faff1ad7ec10/bucket/5b380494-3d7f-4404-913a-310b78f1817f.docx'


def load_template() -> io.BytesIO:
    with urllib.request.urlopen(TEMPLATE_URL, timeout=15) as resp:
        return io.BytesIO(resp.read())

CATEGORY_ORDER = ['cpu', 'motherboard', 'ram', 'gpu', 'ssd', 'cooler', 'psu', 'case']
CATEGORY_LABELS = {
    'cpu': 'Процессор',
    'motherboard': 'Материнская плата',
    'ram': 'Память',
    'gpu': 'Видеокарта',
    'ssd': 'SSD',
    'cooler': 'Охлаждение',
    'psu': 'Блок питания',
    'case': 'Корпус',
}


def check_password(event: dict) -> bool:
    headers = event.get('headers', {}) or {}
    password = headers.get('X-Admin-Password') or headers.get('x-admin-password') or ''
    valid_passwords = [os.environ.get('ADMIN_PASSWORD', ''), os.environ.get('CATALOG_PASSWORD', '')]
    return any(p and hmac.compare_digest(password, p) for p in valid_passwords)


def set_cell_text(cell, text: str):
    p = cell.paragraphs[0]
    if not p.runs:
        p.add_run(text)
        return
    p.runs[0].text = text
    for r in p.runs[1:]:
        r.text = ''


def format_price(value: int) -> str:
    return f'{value:,}'.replace(',', '.')


def build_docx(order: dict, items: list) -> bytes:
    d = docx.Document(load_template())

    p0 = d.paragraphs[0]
    p0.runs[6].text = f"{order['display_number']} "
    date_str = datetime.now().strftime('%d.%m.%y')
    p0.runs[8].text = date_str
    for i in (9, 10, 11, 12, 13):
        p0.runs[i].text = ''

    tbl = d.tables[0]
    rows = list(tbl.rows)
    template_tr = deepcopy(rows[0]._tr)
    total_tr = rows[-1]._tr
    for row in rows[:-1]:
        row._tr.getparent().remove(row._tr)

    sorted_items = sorted(
        items,
        key=lambda it: (
            CATEGORY_ORDER.index(it['category']) if it.get('category') in CATEGORY_ORDER else len(CATEGORY_ORDER),
            it.get('sort_order', 0),
        ),
    )

    for _ in range(len(sorted_items) + 1):
        new_tr = deepcopy(template_tr)
        total_tr.addprevious(new_tr)

    rows2 = list(tbl.rows)
    idx = 1
    for i, it in enumerate(sorted_items):
        label = CATEGORY_LABELS.get(it.get('category'), '')
        text = f"{idx} {label} {it['component_name']}".strip()
        set_cell_text(rows2[i].cells[0], text)
        idx += 1
    set_cell_text(rows2[len(sorted_items)].cells[0], f'{idx} Сборка ПК под ключ')

    set_cell_text(rows2[-1].cells[0], f"Итого {format_price(order['total_price'])} р")

    tbl2 = d.tables[1]
    set_cell_text(tbl2.rows[1].cells[0], f"ФИО {order['customer_name']}")
    set_cell_text(tbl2.rows[1].cells[1], "ФИО Шуляков А.Ю.")
    set_cell_text(tbl2.rows[3].cells[0], "")
    set_cell_text(tbl2.rows[4].cells[0], f"Телефон {order['customer_phone']}")

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def handler(event: dict, context) -> dict:
    """Генерирует гарантийный талон (Word) для заказа CRM и загружает его в S3, возвращая публичную ссылку."""
    cors = {
        'Access-Control-Allow-Origin': '*',
        'Access-Control-Allow-Methods': 'POST, OPTIONS',
        'Access-Control-Allow-Headers': 'Content-Type, X-Admin-Password',
        'Access-Control-Max-Age': '86400',
    }

    if event.get('httpMethod') == 'OPTIONS':
        return {'statusCode': 200, 'headers': cors, 'body': ''}

    if event.get('httpMethod') != 'POST':
        return {'statusCode': 405, 'headers': cors, 'body': json.dumps({'error': 'Метод не поддерживается'})}

    if not check_password(event):
        return {
            'statusCode': 401,
            'headers': {**cors, 'Content-Type': 'application/json'},
            'body': json.dumps({'error': 'Неверный пароль'}),
        }

    body = json.loads(event.get('body') or '{}')
    order_id = body.get('order_id')
    if not order_id:
        return {'statusCode': 400, 'headers': {**cors, 'Content-Type': 'application/json'}, 'body': json.dumps({'error': 'order_id обязателен'})}

    dsn = os.environ['DATABASE_URL']
    conn = psycopg2.connect(dsn)
    cur = conn.cursor()

    try:
        cur.execute(
            f'SELECT id, city, customer_name, customer_phone, final_date, total_price, warranty_number, order_number '
            f'FROM {ORDERS_TABLE} WHERE id = %s',
            (order_id,)
        )
        row = cur.fetchone()
        if not row:
            return {'statusCode': 404, 'headers': {**cors, 'Content-Type': 'application/json'}, 'body': json.dumps({'error': 'Заказ не найден'})}

        order = {
            'id': row[0],
            'city': row[1],
            'customer_name': row[2],
            'customer_phone': row[3],
            'final_date': row[4],
            'total_price': row[5],
            'warranty_number': row[6],
            'order_number': row[7],
        }

        if not order['warranty_number']:
            cur.execute("SELECT nextval('t_p288352_240fps_site_project.warranty_number_seq')")
            num = cur.fetchone()[0]
            order['warranty_number'] = f'{num:03d}'
            cur.execute(
                f'UPDATE {ORDERS_TABLE} SET warranty_number = %s WHERE id = %s',
                (order['warranty_number'], order_id)
            )

        order['display_number'] = order['order_number'] or order['warranty_number']

        cur.execute(
            f'SELECT component_name, category, sort_order FROM {ITEMS_TABLE} WHERE order_id = %s ORDER BY sort_order, id',
            (order_id,)
        )
        items = [{'component_name': r[0], 'category': r[1], 'sort_order': r[2]} for r in cur.fetchall()]

        docx_bytes = build_docx(order, items)

        safe_name = re.sub(r'[^A-Za-z0-9._-]+', '_', order['display_number']).strip('_') or str(order_id)
        file_key = f"warranty/{safe_name}.docx"
        s3 = boto3.client(
            's3',
            endpoint_url='https://bucket.poehali.dev',
            aws_access_key_id=os.environ['AWS_ACCESS_KEY_ID'],
            aws_secret_access_key=os.environ['AWS_SECRET_ACCESS_KEY'],
        )
        s3.put_object(
            Bucket='files',
            Key=file_key,
            Body=docx_bytes,
            ContentType='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            ContentDisposition=f'attachment; filename="{safe_name}.docx"',
        )
        warranty_url = f"https://cdn.poehali.dev/projects/{os.environ['AWS_ACCESS_KEY_ID']}/bucket/{file_key}"

        cur.execute(
            f'UPDATE {ORDERS_TABLE} SET warranty_url = %s WHERE id = %s',
            (warranty_url, order_id)
        )
        conn.commit()

        return {
            'statusCode': 200,
            'headers': {**cors, 'Content-Type': 'application/json'},
            'body': json.dumps({'ok': True, 'url': warranty_url, 'warranty_number': order['warranty_number']}),
        }
    finally:
        cur.close()
        conn.close()